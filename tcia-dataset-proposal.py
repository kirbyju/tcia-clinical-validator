import streamlit as st
import pandas as pd
import os
import io
import json
import zipfile
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from docx import Document
import PyPDF2
import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
import importlib.util
import re
import requests
import uuid

st.set_page_config(page_title="TCIA Dataset Proposal Form", layout="wide")

# Constants & Configuration
RESOURCES_DIR = os.path.join(os.path.dirname(__file__), 'tcia-remapping-skill', 'resources')

# Import MDF parser and helpers
skill_dir = os.path.join(os.path.dirname(__file__), 'tcia-remapping-skill')
mdf_parser_path = os.path.join(skill_dir, 'mdf_parser.py')
spec_mdf = importlib.util.spec_from_file_location("mdf_parser", mdf_parser_path)
mdf_parser = importlib.util.module_from_spec(spec_mdf)
spec_mdf.loader.exec_module(mdf_parser)
get_mdf_resources = mdf_parser.get_mdf_resources

orcid_helper_path = os.path.join(skill_dir, 'orcid_helper.py')
spec_orcid = importlib.util.spec_from_file_location("orcid_helper", orcid_helper_path)
orcid_helper = importlib.util.module_from_spec(spec_orcid)
spec_orcid.loader.exec_module(orcid_helper)

PERMISSIBLE_VALUES_FILE = os.path.join(RESOURCES_DIR, 'permissible_values.json')

LABELS = {
    "Proposal Type": "What kind of dataset are you submitting?",
    "Scientific POC Name": "Scientific POC Name*",
    "Scientific POC Email": "Scientific POC Email*",
    "Technical POC Name": "Technical POC Name*",
    "Technical POC Email": "Technical POC Email*",
    "Legal POC Name": "Legal POC Name*",
    "Legal POC Email": "Legal POC Email*",
    "Title": "Suggest a descriptive title for your dataset*",
    "Nickname": "Suggest a shorter nickname for your dataset*",
    "Authors": "List the authors of this data set*",
    "Abstract": "Dataset Abstract*",
    "Published Elsewhere": "Has this data ever been published elsewhere?*",
    "disease_site": "Primary disease site/location*",
    "diagnosis": "Histologic diagnosis*",
    "image_types": "Which image types are included in the data set?*",
    "supporting_data": "Which kinds of supporting data are included in the data set?*",
    "file_formats": "Specify the file format utilized for each type of data*",
    "modifications": "Describe any steps taken to de-identify or otherwise modify the data in preparation for submission*",
    "faces": "Does your data contain any images of patient faces?*",
    "collections_analyzed": "Which TCIA collection(s) did you analyze?*",
    "derived_types": "What types of derived data are included in the dataset?*",
    "image_records": "Do you have records to indicate exactly which TCIA images analyzed?*",
    "disk_space": "Approximate disk space required*",
    "Time Constraints": "Are there any time constraints associated with sharing your data set?*",
    "descriptor_publication": "Is there a related dataset descriptor publication? (i.e. Nature Scientific Data article on how to use the dataset)*",
    "additional_publications": "Any additional publications derived from these data?*",
    "acknowledgments": "Acknowledgments or funding statements*",
    "why_tcia": "Why would you like to publish this dataset on TCIA?*"
}

IMAGE_FORMATS = [
    "DICOM", "NIfTI", "PNG", "JPEG", "Aperio (.svs, .tif)",
    "Hamamatsu (.vms, .vmu, .ndpi)", "Leica (.scn)", "MIRAX (.mrxs)",
    "Philips (.tiff)", "Sakura (.svslide)", "Trestle (.tif)",
    "Ventana (.bif, .tif)", "Generic tiled TIFF (.tif)", "Other"
]
SUPPORT_FORMATS = ["Tabular (.csv, .tsv, .xlsx)", "JSON", "Text", "Other"]
COMBINED_FORMATS = [
    "DICOM", "NIfTI", "PNG", "JPEG", "Aperio (.svs, .tif)",
    "Hamamatsu (.vms, .vmu, .ndpi)", "Leica (.scn)", "MIRAX (.mrxs)",
    "Philips (.tiff)", "Sakura (.svslide)", "Trestle (.tif)",
    "Ventana (.bif, .tif)", "Generic tiled TIFF (.tif)",
    "Tabular (.csv, .tsv, .xlsx)", "JSON", "Text", "Other"
]

def load_json(filepath):
    with open(filepath, 'r') as f:
        return json.load(f)

def lookup_doi(doi):
    """Fetch metadata from Crossref API"""
    if not doi:
        return None
    try:
        url = f"https://api.crossref.org/works/{doi}"
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()['message']
            title = data.get('title', [''])[0]
            authors_list = data.get('author', [])
            authors = ", ".join([f"{a.get('family', '')} {a.get('given', '')}".strip() for a in authors_list])
            if len(authors_list) > 3:
                authors = f"{authors_list[0].get('family', '')} et al."
            year = ""
            issued = data.get('issued', {}).get('date-parts', [[None]])[0][0]
            if issued:
                year = str(issued)
            journal = data.get('container-title', [''])[0]
            return f"{authors} ({year}). {title}. {journal}. DOI: {doi}"
    except:
        pass
    return None

@st.cache_data
def load_mdf_data():
    schema, mdf_pv, relationships = get_mdf_resources(RESOURCES_DIR)
    legacy_pv = load_json(PERMISSIBLE_VALUES_FILE)

    final_pv = legacy_pv.copy()
    if mdf_pv:
        for k, v in mdf_pv.items():
            final_pv[k] = v
    return schema, final_pv, relationships

schema, permissible_values, relationships = load_mdf_data()
AGREEMENT_TEMPLATE = os.path.join(RESOURCES_DIR, 'agreement_template.pdf')
HELP_DESK_EMAIL = os.getenv("TCIA_HELP_DESK_EMAIL", "help@cancerimagingarchive.net")

# Server-side SMTP configuration
SMTP_SERVER = os.getenv("SMTP_SERVER")
SMTP_PORT = int(os.getenv("SMTP_PORT", 587))
SMTP_USER = os.getenv("SMTP_USER")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")

st.markdown(
    """
    <div style="display: flex; align-items: center; margin-bottom: 20px;">
        <img src="https://www.cancerimagingarchive.net/wp-content/uploads/2021/06/TCIA-Logo-01.png" style="height: 50px; margin-right: 15px;">
        <h1 style="margin: 0;">Dataset Proposal Form</h1>
    </div>
    """,
    unsafe_allow_html=True
)


# Initial choice
proposal_type = st.radio(
    "What kind of dataset are you submitting?",
    options=["New Collection Proposal", "Analysis Results Proposal"],
    help="Select 'New Collection' if you are submitting primary imaging data. Select 'Analysis Results' if you are submitting derived data (e.g., segmentations) from existing TCIA collections.",
    key="prop_type"
)

if proposal_type == "New Collection Proposal":
    st.info("💡 **New Collection Proposal**: Used for contributing original imaging data that has not been previously hosted on TCIA.")
else:
    st.info("💡 **Analysis Results Proposal**: Used for contributing derived data like annotations, segmentations, or radiomics features based on existing TCIA datasets.")

st.markdown("---")

# Form Questions
st.subheader("Contact Information")
if SMTP_SERVER and SMTP_USER and SMTP_PASSWORD:
    st.info("💡 **Note**: After completing this form, a copy of this proposal will be emailed to the Scientific, Technical, and Legal points of contact listed below.")

col1, col2 = st.columns(2)
with col1:
    sci_poc_name = st.text_input(LABELS["Scientific POC Name"], help="The person to contact about the proposal and data collection.", key="sci_poc_name")
    tech_poc_name = st.text_input(LABELS["Technical POC Name"], help="The person involved in sending the data.", key="tech_poc_name")
    legal_poc_name = st.text_input(LABELS["Legal POC Name"], help="Authorized signatory who will sign the TCIA Data Submission Agreement. This should not be the PI or department chair.", key="legal_poc_name")
with col2:
    sci_poc_email = st.text_input(LABELS["Scientific POC Email"], key="sci_poc_email")
    tech_poc_email = st.text_input(LABELS["Technical POC Email"], key="tech_poc_email")
    legal_poc_email = st.text_input(LABELS["Legal POC Email"], key="legal_poc_email")

st.subheader("Dataset Publication Details")
title = st.text_input(LABELS["Title"], help="Similar to a manuscript title.", key="title")
nickname = st.text_input(LABELS["Nickname"], help="Must be < 30 characters, alphanumeric and dashes only.", max_chars=30, key="nickname")
if nickname and not re.match(r"^[a-zA-Z0-9-]+$", nickname):
    st.error("⚠️ Invalid nickname. Only letters, numbers, and dashes are allowed.")

# Authors Section
st.write(f"**{LABELS['Authors']}**")
st.markdown("""
Please list all authors in the order you'd like them to appear if the data set is published on TCIA.
Names should be listed as **(FAMILY, GIVEN) - ORCID**.
You can quickly register for an ORCID at [https://orcid.org/signin](https://orcid.org/signin) if you don't have one already.
""")
authors_raw = st.text_area("Example: Smith, John - 0000-0002-1234-5678; Doe, Jane",
                           help="List authors one per line or separated by semicolons.",
                           key="authors_raw_input")

abstract = st.text_area(LABELS["Abstract"], help="Focus on describing the dataset itself.", max_chars=1000, key="abstract")

st.subheader("Data Collection Details")
published_elsewhere = st.text_input(LABELS["Published Elsewhere"], help="If so, why publish on TCIA? Do you intend for the original to remain accessible?", key="published_elsewhere")

# Adaptive fields
extra_data = {}
if proposal_type == "New Collection Proposal":
    col1, col2 = st.columns(2)
    with col1:
        site_raw = permissible_values.get('primary_site', []) if permissible_values else []
        site_options = sorted(list(set([v['value'] if isinstance(v, dict) else str(v) for v in site_raw])))
        extra_data['disease_site'] = st.selectbox(LABELS["disease_site"], options=[""] + site_options, key="disease_site")
    with col2:
        diag_raw = permissible_values.get('primary_diagnosis', []) if permissible_values else []
        diag_options = sorted(list(set([v['value'] if isinstance(v, dict) else str(v) for v in diag_raw])))
        extra_data['diagnosis'] = st.selectbox(LABELS["diagnosis"], options=[""] + diag_options, key="diagnosis")

    extra_data['image_types'] = st.multiselect(
        LABELS["image_types"],
        options=["MR", "CT", "PET", "Mammograms", "Ultrasound", "Xray", "Radiation Therapy", "Whole Slide Image", "CODEX", "Single-cell Image", "Photomicrograph", "Microarray", "Multiphoton", "Immunofluorescence", "Other"],
        key="image_types"
    )
    extra_data['supporting_data'] = st.multiselect(
        LABELS["supporting_data"],
        options=["Clinical", "Image Analyses", "Image Registrations", "Genomics", "Proteomics", "Software / Source Code", "No additional data", "Other"],
        key="supporting_data"
    )

    st.write(f"**{LABELS['file_formats']}**")
    nc_formats = []
    for itype in extra_data.get('image_types', []):
        col1, col2 = st.columns([1, 1])
        with col1: st.write(f"Format for **{itype}**:")
        with col2:
            fmt = st.selectbox(f"Select format for {itype}", options=IMAGE_FORMATS, key=f"fmt_img_{itype}", label_visibility="collapsed")
            if fmt == "Other":
                other_fmt = st.text_input(f"Specify other format for {itype}", key=f"other_fmt_img_{itype}", label_visibility="collapsed")
                nc_formats.append(f"{itype} - {other_fmt}" if other_fmt else f"{itype} - Other")
            else:
                nc_formats.append(f"{itype} - {fmt}")
    for stype in extra_data.get('supporting_data', []):
        if stype == "No additional data": continue
        col1, col2 = st.columns([1, 1])
        with col1: st.write(f"Format for **{stype}**:")
        with col2:
            opts = COMBINED_FORMATS if stype == "Image Analyses" else SUPPORT_FORMATS
            fmt = st.selectbox(f"Select format for {stype}", options=opts, key=f"fmt_supp_{stype}", label_visibility="collapsed")
            if fmt == "Other":
                other_fmt = st.text_input(f"Specify other format for {stype}", key=f"other_fmt_supp_{stype}", label_visibility="collapsed")
                nc_formats.append(f"{stype} - {other_fmt}" if other_fmt else f"{stype} - Other")
            else:
                nc_formats.append(f"{stype} - {fmt}")
    extra_data['file_formats'] = "; ".join(nc_formats)

    extra_data['modifications'] = st.text_area(LABELS["modifications"], key="modifications")
    extra_data['faces'] = st.radio(LABELS["faces"], options=["Yes", "No"], key="faces")
else: # Analysis Results
    extra_data['collections_analyzed'] = st.text_input(LABELS["collections_analyzed"], key="collections_analyzed")
    extra_data['derived_types'] = st.multiselect(
        LABELS["derived_types"],
        options=["Segmentation", "Classification", "Quantitative Feature", "Image (converted/processed/registered)", "Other"],
        key="derived_types"
    )
    extra_data['image_records'] = st.radio(LABELS["image_records"], options=["Yes, I know exactly.", "No, I need assistance."], key="image_records")

    st.write(f"**{LABELS['file_formats']}**")
    ar_formats = []
    for dtype in extra_data.get('derived_types', []):
        col1, col2 = st.columns([1, 1])
        with col1: st.write(f"Format for **{dtype}**:")
        with col2:
            fmt = st.selectbox(f"Select format for {dtype}", options=COMBINED_FORMATS, key=f"fmt_ar_{dtype}", label_visibility="collapsed")
            if fmt == "Other":
                other_fmt = st.text_input(f"Specify other format for {dtype}", key=f"other_fmt_ar_{dtype}", label_visibility="collapsed")
                ar_formats.append(f"{dtype} - {other_fmt}" if other_fmt else f"{dtype} - Other")
            else:
                ar_formats.append(f"{dtype} - {fmt}")
    extra_data['file_formats'] = "; ".join(ar_formats)
# Shared bottom fields
col1, col2 = st.columns(2)
with col1:
    extra_data['disk_space'] = st.text_input(LABELS["disk_space"], key="disk_space")
with col2:
    time_constraints = st.text_input(LABELS["Time Constraints"], key="time_constraints")

extra_data['descriptor_publication'] = st.text_area(LABELS['descriptor_publication'], key="descriptor_publication")
extra_data['additional_publications'] = st.text_area(LABELS['additional_publications'], key="additional_publications")
extra_data['acknowledgments'] = st.text_area(LABELS["acknowledgments"], key="acknowledgments")
extra_data['why_tcia'] = st.multiselect(
    LABELS["why_tcia"],
    options=["To meet a funding agency's requirements", "To meet a journal's requirements", "To facilitate collaboration", "To facilitate a challenge competition", "Other"],
    key="why_tcia"
)

button_label = "Submit" if (SMTP_SERVER and SMTP_USER and SMTP_PASSWORD) else "Generate Proposal Documents"
submit_button = st.button(button_label, type="primary")

# Processing after submission
if submit_button:
    # 1. Validation
    missing_fields = []
    if not sci_poc_name: missing_fields.append(LABELS["Scientific POC Name"])
    if not sci_poc_email: missing_fields.append(LABELS["Scientific POC Email"])
    if not tech_poc_name: missing_fields.append(LABELS["Technical POC Name"])
    if not tech_poc_email: missing_fields.append(LABELS["Technical POC Email"])
    if not legal_poc_name: missing_fields.append(LABELS["Legal POC Name"])
    if not legal_poc_email: missing_fields.append(LABELS["Legal POC Email"])
    if not title: missing_fields.append(LABELS["Title"])
    if not nickname:
        missing_fields.append(LABELS["Nickname"])
    elif not re.match(r"^[a-zA-Z0-9-]+$", nickname):
        st.error("⚠️ Invalid nickname. Please use only letters, numbers, and dashes.")
        st.stop()
    if not authors_raw:
        missing_fields.append(LABELS["Authors"])
    if not abstract: missing_fields.append(LABELS["Abstract"])
    if not published_elsewhere: missing_fields.append(LABELS["Published Elsewhere"])

    for key, val in extra_data.items():
        if not val:
            label = LABELS.get(key, key.replace('_', ' ').title())
            missing_fields.append(label)

    if not time_constraints: missing_fields.append(LABELS["Time Constraints"])

    if missing_fields:
        st.error("Please fill in all required fields:")
        for field in missing_fields:
            st.write(f"- {field}")
    else:
        # Prepare data for files
        all_responses = {
            "Proposal Type": proposal_type,
            "Scientific POC Name": sci_poc_name,
            "Scientific POC Email": sci_poc_email,
            "Technical POC Name": tech_poc_name,
            "Technical POC Email": tech_poc_email,
            "Legal POC Name": legal_poc_name,
            "Legal POC Email": legal_poc_email,
            "Time Constraints": time_constraints,
            "Title": title,
            "Nickname": nickname,
            "Authors": authors_raw,
            "Abstract": abstract,
            "Published Elsewhere": published_elsewhere
        }
        all_responses.update(extra_data)

        # Generate TSV
        tsv_buffer = io.StringIO()
        df = pd.DataFrame([all_responses])
        df.to_csv(tsv_buffer, sep='\t', index=False)

        # Generate DOCX
        doc = Document()
        doc.add_heading(f"TCIA Dataset Proposal: {title}", 0)
        for key, value in all_responses.items():
            label = LABELS.get(key, key)

            doc.add_paragraph(f"{label}:", style='Heading 2')
            doc.add_paragraph(str(value))

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Generate PDF
        pdf_buffer = io.BytesIO()
        try:
            # Create a new page 6 for Exhibit A
            exhibit_a_buffer = io.BytesIO()
            c = canvas.Canvas(exhibit_a_buffer, pagesize=letter)
            width, height = letter

            # Header matching the template
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 50, "TCIA Data Submission Agreement (v. 20220 914) Page 6 of 7")

            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width/2, height - 100, "EXHIBIT A")
            c.drawCentredString(width/2, height - 120, "DESCRIPTION OF SUBMISSION DATA")

            c.setFont("Helvetica", 11)
            text_object = c.beginText(50, height - 160)
            text_object.setFont("Helvetica", 11)

            wrapped_abstract = simpleSplit(abstract, "Helvetica", 11, width - 100)
            for line in wrapped_abstract:
                text_object.textLine(line)

            c.drawText(text_object)
            c.showPage()
            c.save()
            exhibit_a_buffer.seek(0)

            # Merge with template
            reader = PyPDF2.PdfReader(AGREEMENT_TEMPLATE)
            writer = PyPDF2.PdfWriter()

            # Pages 1-5
            for i in range(5):
                writer.add_page(reader.pages[i])

            # Replace Page 6
            new_exhibit_reader = PyPDF2.PdfReader(exhibit_a_buffer)
            writer.add_page(new_exhibit_reader.pages[0])

            # Page 7
            writer.add_page(reader.pages[6])

            writer.write(pdf_buffer)
            pdf_buffer.seek(0)
            pdf_success = True
        except Exception as e:
            st.error(f"Error generating PDF: {e}")
            pdf_success = False

        # Create ZIP
        today = datetime.date.today().isoformat()
        pkg_name = f"{nickname}_proposal_package_{today}.zip"
        summary_tsv_name = f"{nickname}_proposal_summary_{today}.tsv"
        docx_name = f"{nickname}_proposal_summary_{today}.docx"
        pdf_name = f"{nickname}_agreement_updated_{today}.pdf"

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
            zip_file.writestr(summary_tsv_name, tsv_buffer.getvalue())
            zip_file.writestr(docx_name, docx_buffer.getvalue())
            if pdf_success:
                zip_file.writestr(pdf_name, pdf_buffer.getvalue())
        zip_buffer.seek(0)

        # Store in session state for later use in email and to persist buttons
        st.session_state['proposal_files'] = {
            'tsv': tsv_buffer.getvalue(),
            'docx': docx_buffer.getvalue(),
            'pdf': pdf_buffer.getvalue() if pdf_success else None,
            'zip': zip_buffer.getvalue(),
            'title': title,
            'nickname': nickname,
            'proposal_type': proposal_type,
            'pocs': [sci_poc_email, tech_poc_email, legal_poc_email],
            'pkg_name': pkg_name
        }
        st.session_state['proposal_generated'] = True

        # Determine if we can send automatically
        can_send_auto = all([SMTP_SERVER, SMTP_USER, SMTP_PASSWORD])
        if can_send_auto:
            try:
                msg = MIMEMultipart()
                msg['From'] = SMTP_USER
                msg['To'] = HELP_DESK_EMAIL
                msg['Cc'] = ", ".join([sci_poc_email, tech_poc_email, legal_poc_email])
                msg['Subject'] = f"New Dataset Proposal: {title}"

                body = f"A new dataset proposal has been submitted via the Dataset Proposal Form.\n\nType: {proposal_type}\nTitle: {title}\nSubmitters: {', '.join([sci_poc_email, tech_poc_email, legal_poc_email])}"
                msg.attach(MIMEText(body, 'plain'))

                part = MIMEBase('application', 'zip')
                part.set_payload(zip_buffer.getvalue())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{pkg_name}"')
                msg.attach(part)

                server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
                server.starttls()
                server.login(SMTP_USER, SMTP_PASSWORD)
                # Send to both To and Cc addresses
                recipients = [HELP_DESK_EMAIL, sci_poc_email, tech_poc_email, legal_poc_email]
                server.send_message(msg, to_addrs=recipients)
                server.quit()
                st.session_state['email_sent'] = True
            except Exception as e:
                st.error(f"Failed to send email: {e}")

if st.session_state.get('proposal_generated'):
    can_send_auto = all([SMTP_SERVER, SMTP_USER, SMTP_PASSWORD])
    files = st.session_state.get('proposal_files')

    if not can_send_auto:
        st.success("✅ Proposal documents generated successfully!")

        col1, col2, col3 = st.columns(3)
        today = datetime.date.today().isoformat()
        nickname_val = files.get('nickname', 'dataset')
        with col1:
            st.download_button("Download TSV (Import to Remapper)", data=files['tsv'], file_name=f"{nickname_val}_proposal_{today}.tsv", mime="text/tab-separated-values")
        with col2:
            st.download_button("Download DOCX Summary", data=files['docx'], file_name=f"{nickname_val}_proposal_summary_{today}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col3:
            if files['pdf']:
                st.download_button("Download PDF Agreement", data=files['pdf'], file_name=f"{nickname_val}_agreement_updated_{today}.pdf", mime="application/pdf")

        st.info("💡 **Please download these documents and keep a copy for your records.**")
        st.info(f"💡 **Manual Submission**: Please email the generated files as attachments to **{HELP_DESK_EMAIL}**.")
    else:
        if st.session_state.get('email_sent'):
            st.success(f"📨 Proposal sent to {HELP_DESK_EMAIL} and CC'd to POCs!")
        else:
            st.info("Proposal generated but failed to send email. Please check your SMTP settings or contact support.")

# Footer
st.markdown("---")
st.markdown(f"""
<div style='text-align: center; color: gray; font-size: 0.9em;'>
TCIA Dataset Proposal Form | {HELP_DESK_EMAIL}
</div>
""", unsafe_allow_html=True)
